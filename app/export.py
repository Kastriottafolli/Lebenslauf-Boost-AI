"""Export des Lebenslaufs als PDF (ReportLab) oder Word/DOCX (python-docx).

Drei professionelle Designs ('classic', 'modern', 'minimal') mit optionalem
Bewerbungsfoto im Kopfbereich. Der vom Modell erzeugte Markdown-Lebenslauf wird
geparst und layoutet.
"""

import html
import io
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# Farb- und Stilpalette je Design.
_PALETTE = {
    "classic": {"accent": "#1A1A1A", "muted": "#444444", "rule": "#1A1A1A", "serif": True},
    "modern": {"accent": "#A06A24", "muted": "#5A5248", "rule": "#A06A24", "serif": False},
    "minimal": {"accent": "#222222", "muted": "#777777", "rule": "#CCCCCC", "serif": False},
}

_HEADER_TYPES = ("name", "title", "contact")


# ──────────────────────────────────────────────────────────────────────────
# Markdown-Parser
# ──────────────────────────────────────────────────────────────────────────
def parse_cv(md: str) -> List[dict]:
    elements: List[dict] = []
    pending_sub: List[str] = []
    seen_section = False

    def flush_sub() -> None:
        nonlocal pending_sub
        if pending_sub:
            elements.append({"type": "title", "text": pending_sub[0]})
            if len(pending_sub) > 1:
                elements.append({"type": "contact", "text": "  ".join(pending_sub[1:])})
            pending_sub = []

    for raw in md.splitlines():
        s = _clean(raw.strip())
        if not s:
            continue
        if s.startswith("> "):
            flush_sub()
            elements.append({"type": "note", "text": s[2:].strip()})
        elif s.startswith("# "):
            flush_sub()
            elements.append({"type": "name", "text": s[2:].strip()})
        elif s.startswith("### "):
            flush_sub()
            seen_section = True
            elements.append({"type": "entry", "text": s[4:].strip()})
        elif s.startswith("## "):
            flush_sub()
            seen_section = True
            elements.append({"type": "section", "text": s[3:].strip()})
        elif s.startswith("- ") or s.startswith("* "):
            flush_sub()
            elements.append({"type": "bullet", "text": s[2:].strip()})
        else:
            if (elements and elements[-1]["type"] == "name") or (
                pending_sub and not seen_section
            ):
                pending_sub.append(s)
            elif not seen_section and elements and elements[-1]["type"] in ("title", "contact"):
                pending_sub.append(s)
            else:
                elements.append({"type": "para", "text": s})

    flush_sub()
    return elements


def _clean(s: str) -> str:
    return s.replace("**", "").replace("`", "").replace("__", "")


# ══════════════════════════════════════════════════════════════════════════
# PDF (ReportLab)
# ══════════════════════════════════════════════════════════════════════════
def to_pdf(
    content: str, design: str = "modern", language: str = "de", photo: Optional[bytes] = None
) -> bytes:
    pal = _PALETTE.get(design, _PALETTE["modern"])
    serif = pal["serif"]
    base = "Times-Roman" if serif else "Helvetica"
    bold = "Times-Bold" if serif else "Helvetica-Bold"
    italic = "Times-Italic" if serif else "Helvetica-Oblique"
    accent = HexColor(pal["accent"])
    muted = HexColor(pal["muted"])
    rule = HexColor(pal["rule"])
    name_align = TA_CENTER if design == "classic" else TA_LEFT

    styles = {
        "name": ParagraphStyle(
            "name", fontName=bold, fontSize=23 if design == "minimal" else 25,
            textColor=HexColor("#111111"), alignment=name_align, spaceAfter=2, leading=27,
        ),
        "title": ParagraphStyle(
            "title", fontName=italic if design == "classic" else base, fontSize=12.5,
            textColor=accent if design == "modern" else muted, alignment=name_align, spaceAfter=3,
        ),
        "contact": ParagraphStyle(
            "contact", fontName=base, fontSize=9, textColor=muted,
            alignment=name_align, spaceAfter=2, leading=12,
        ),
        "section": ParagraphStyle(
            "section", fontName=bold, fontSize=12.5, textColor=accent,
            spaceBefore=13, spaceAfter=2,
        ),
        "entry": ParagraphStyle(
            "entry", fontName=bold, fontSize=10.5, textColor=HexColor("#222222"),
            spaceBefore=6, spaceAfter=1,
        ),
        "bullet": ParagraphStyle(
            "bullet", fontName=base, fontSize=10, textColor=HexColor("#333333"),
            leftIndent=12, spaceAfter=2, leading=13.5,
        ),
        "para": ParagraphStyle(
            "para", fontName=base, fontSize=10, textColor=HexColor("#333333"),
            spaceAfter=3, leading=13.5,
        ),
        "note": ParagraphStyle(
            "note", fontName=base, fontSize=8.5, textColor=HexColor("#B45309"), spaceAfter=8,
        ),
    }

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm, title="Lebenslauf",
    )
    els = parse_cv(content)
    story: list = []

    # Demo-/Hinweis-Zeilen ganz oben.
    for e in els:
        if e["type"] == "note":
            story.append(Paragraph(_xml(e["text"]), styles["note"]))

    # Akzentleiste oben (nur Modern).
    if design == "modern":
        story.append(HRFlowable(width="100%", thickness=3, color=accent, spaceAfter=8))

    # Kopfbereich (Name/Titel/Kontakt + Foto).
    header_els = [e for e in els if e["type"] in _HEADER_TYPES]
    story.extend(_pdf_header(header_els, photo, design, styles, doc.width))

    # Hauptteil.
    for e in els:
        t, text = e["type"], _xml(e["text"])
        if t in _HEADER_TYPES or t == "note":
            continue
        if t == "section":
            label = text.upper() if design in ("classic", "minimal") else text
            story.append(Paragraph(label, styles["section"]))
            thickness = 0.5 if design == "minimal" else (1.5 if design == "modern" else 1)
            story.append(HRFlowable(width="100%", thickness=thickness, color=rule,
                                    spaceBefore=1, spaceAfter=4))
        elif t == "entry":
            story.append(Paragraph(text, styles["entry"]))
        elif t == "bullet":
            marker = "–&nbsp;&nbsp;" if design == "classic" else "•&nbsp;&nbsp;"
            story.append(Paragraph(marker + text, styles["bullet"]))
        else:
            story.append(Paragraph(text, styles["para"]))

    if len(story) <= 1:
        story.append(Spacer(1, 1))
    doc.build(story)
    return buf.getvalue()


def _pdf_header(header_els, photo, design, styles, avail_width) -> list:
    text_cells = [Paragraph(_xml(e["text"]), styles[e["type"]]) for e in header_els]
    if not text_cells:
        text_cells = [Spacer(1, 1)]

    img = None
    if photo:
        try:
            sizes = {"classic": (2.7, 3.4), "minimal": (2.4, 3.0), "modern": (2.8, 3.5)}
            w, h = sizes.get(design, sizes["modern"])
            img = RLImage(io.BytesIO(photo), width=w * cm, height=h * cm)
        except Exception:
            img = None

    if img is None:
        return [*text_cells, Spacer(1, 6)]

    if design == "classic":
        img.hAlign = "CENTER"
        return [img, Spacer(1, 6), *text_cells, Spacer(1, 8)]

    photo_w = (2.4 if design == "minimal" else 2.8) * cm + 0.4 * cm
    table = Table([[text_cells, img]], colWidths=[avail_width - photo_w, photo_w])
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    return [table, Spacer(1, 10)]


def _xml(s: str) -> str:
    return html.escape(s, quote=False)


# ══════════════════════════════════════════════════════════════════════════
# Word / DOCX (python-docx)
# ══════════════════════════════════════════════════════════════════════════
def to_docx(
    content: str, design: str = "modern", language: str = "de", photo: Optional[bytes] = None
) -> bytes:
    pal = _PALETTE.get(design, _PALETTE["modern"])
    font_name = "Georgia" if pal["serif"] else "Calibri"
    accent = pal["accent"].lstrip("#")
    muted = pal["muted"].lstrip("#")
    rule = pal["rule"].lstrip("#")
    center = design == "classic"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(10.5)

    els = parse_cv(content)

    for e in els:
        if e["type"] == "note":
            p = doc.add_paragraph()
            r = p.add_run(e["text"])
            r.italic = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor.from_string("B45309")

    header_els = [e for e in els if e["type"] in _HEADER_TYPES]
    _docx_header(doc, header_els, photo, design, accent, muted, center)

    for e in els:
        t, text = e["type"], e["text"]
        if t in _HEADER_TYPES or t == "note":
            continue
        if t == "section":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(11)
            run = p.add_run(text.upper() if design in ("classic", "minimal") else text)
            run.bold = True
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor.from_string(accent)
            _bottom_border(p, color=rule, size=12 if design == "modern" else 6)
        elif t == "entry":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string("222222")
        elif t == "bullet":
            p = doc.add_paragraph(text, style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_line(paragraph, text, kind, accent, muted, center):
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    if kind == "name":
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor.from_string("111111")
    elif kind == "title":
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor.from_string(accent)
    else:  # contact
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(muted)


def _docx_header(doc, header_els, photo, design, accent, muted, center) -> None:
    if not photo:
        for e in header_els:
            _docx_line(doc.add_paragraph(), e["text"], e["type"], accent, muted, center)
        return

    if design == "classic":
        pic_par = doc.add_paragraph()
        pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_par.add_run().add_picture(io.BytesIO(photo), width=Cm(2.7))
        for e in header_els:
            _docx_line(doc.add_paragraph(), e["text"], e["type"], accent, muted, True)
        return

    # Modern/Minimal: zweispaltige, randlose Tabelle (Text links, Foto rechts).
    photo_w = 2.4 if design == "minimal" else 2.8
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    left, right = table.rows[0].cells
    left.width = Cm(13.5)
    right.width = Cm(photo_w + 0.3)

    left.paragraphs[0]._p.getparent().remove(left.paragraphs[0]._p)
    for e in header_els:
        _docx_line(left.add_paragraph(), e["text"], e["type"], accent, muted, False)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run().add_picture(io.BytesIO(photo), width=Cm(photo_w))


def _bottom_border(paragraph, color: str = "000000", size: int = 6) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)
