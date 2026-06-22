"""Text- und Foto-Extraktion aus hochgeladenen Lebensläufen (PDF / DOCX / TXT)."""

import io
from typing import List, Optional, Tuple

from docx import Document
from pypdf import PdfReader


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text = _extract_pdf(data)
    elif name.endswith(".docx"):
        text = _extract_docx(data)
    elif name.endswith(".txt"):
        text = data.decode("utf-8", errors="ignore")
    elif name.endswith(".doc"):
        raise ValueError(
            "Das alte .doc-Format wird nicht unterstützt. "
            "Bitte als .docx oder PDF speichern."
        )
    else:
        raise ValueError("Nicht unterstütztes Format. Bitte PDF, DOCX oder TXT nutzen.")

    text = text.strip()
    if not text:
        raise ValueError(
            "Es konnte kein Text aus der Datei gelesen werden "
            "(evtl. ein gescanntes Bild-PDF?)."
        )
    return text


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ──────────────────────────────────────────────────────────────────────────
# Foto-Extraktion: das Bewerbungsfoto aus dem hochgeladenen Lebenslauf holen.
# Heuristik: hochformatiges/quadratisches Bild mit hoher Bytes-pro-Pixel-Dichte
# (echtes Foto), keine flachen Hintergrund-/Banner-Bilder.
# ──────────────────────────────────────────────────────────────────────────
def extract_photo(filename: str, data: bytes) -> Optional[bytes]:
    """Gibt ein normalisiertes Portrait-JPEG zurück, oder None."""
    name = (filename or "").lower()
    raw_images: List[bytes] = []
    if name.endswith(".pdf"):
        raw_images = _pdf_images(data)
    elif name.endswith(".docx"):
        raw_images = _docx_images(data)
    else:
        return None

    best = _pick_best_photo(raw_images)
    return _normalize_photo(best) if best else None


def _pdf_images(data: bytes) -> List[bytes]:
    out: List[bytes] = []
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception:
        return out
    for page in reader.pages:
        try:
            for img in page.images:
                if img.data:
                    out.append(img.data)
        except Exception:
            continue
    return out


def _docx_images(data: bytes) -> List[bytes]:
    out: List[bytes] = []
    try:
        doc = Document(io.BytesIO(data))
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                out.append(rel.target_part.blob)
    except Exception:
        pass
    return out


def _pick_best_photo(images: List[bytes]) -> Optional[bytes]:
    from PIL import Image

    candidates: List[Tuple[float, bytes]] = []
    for blob in images:
        try:
            im = Image.open(io.BytesIO(blob))
            w, h = im.size
        except Exception:
            continue
        if w < 80 or h < 80:
            continue  # zu klein -> Icon/Logo
        ratio = w / h
        if ratio < 0.5 or ratio > 1.4:
            continue  # zu breit -> Banner/Hintergrund
        bpp = len(blob) / (w * h)  # Bytes pro Pixel ~ Detailgrad
        if bpp < 0.05:
            continue  # flacher Hintergrund/Verlauf
        portrait_bonus = 1.0 if h >= w else 0.75
        candidates.append((bpp * portrait_bonus, blob))

    if not candidates:
        return None
    candidates.sort(key=lambda c: c[0], reverse=True)
    return candidates[0][1]


def _normalize_photo(
    blob: bytes, target_ratio: float = 0.8, max_w: int = 500
) -> Optional[bytes]:
    """Zentriert auf Portrait-Format (4:5) zuschneiden, als JPEG ausgeben."""
    from PIL import Image

    try:
        im = Image.open(io.BytesIO(blob)).convert("RGB")
    except Exception:
        return None
    w, h = im.size
    if w / h > target_ratio:  # zu breit -> Seiten beschneiden
        new_w = int(h * target_ratio)
        x = (w - new_w) // 2
        im = im.crop((x, 0, x + new_w, h))
    else:  # zu hoch -> oben/unten beschneiden
        new_h = int(w / target_ratio)
        y = (h - new_h) // 2
        im = im.crop((0, y, w, y + new_h))
    if im.width > max_w:
        im = im.resize((max_w, int(max_w / target_ratio)))
    out = io.BytesIO()
    im.save(out, "JPEG", quality=88)
    return out.getvalue()
